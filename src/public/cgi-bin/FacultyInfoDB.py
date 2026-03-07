import mysql.connector
from docx import Document

#env testing
import os
from dotenv import load_dotenv, dotenv_values
print("hi")

load_dotenv()

"""
you can test out this open source database in case the cpanel database isn't working
HOST=mysql-rfam-public.ebi.ac.uk
ENVUSER=rfamro
PASSWORD=
PORT=4497
NAME=Rfam
"""
doc = Document()
table = doc.add_table(rows=1, cols=12)

table.style = 'Table Grid'
table.cell(0, 0).text = "Faculty Name"
table.cell(0, 1).text = "Highest Degree Earned- Field and Year"
table.cell(0, 2).text = "Rank 1"
table.cell(0, 3).text = "Type of Academic Appointment"
table.cell(0, 4).text = "Govt./Ind. Practice"
table.cell(0, 5).text = "Teaching"
table.cell(0, 6).text = "This Institution"
table.cell(0, 7).text = "Professional Registration/ Certification"
table.cell(0, 8).text = "Professional Organizations"
table.cell(0, 9).text = "This Institution"
table.cell(0, 10).text = "Professional Development"
table.cell(0, 11).text = "Consulting/summer work in industry"





nummy = 0
#for i in dict_output:
#     row[nummy].text = str(dict_output[i])
#     nummy += 1


        #maybe?
        # host=os.getenv("MYSQL_HOSTNAME"),
        # user=os.getenv("MYSQL_USER"),
        # password=os.getenv("MYSQL_PASS"),
        # port=os.getenv("MYSQL_PORT"),
        # database=os.getenv("MYSQL_DATABASE")
        # also change table name

try:
    conn = mysql.connector.connect(
        host="mysql-rfam-public.ebi.ac.uk",
        user="rfamro",
        password="",
        port=4497,
        database="Rfam"
    )

    curr = conn.cursor()
    curr.execute("SELECT * From full_region limit 10;")

    for i in curr.fetchall():
        #creates tables here
        row = table.add_row().cells

        for j in range(len(i)):
            row[j].text = str(i[j])
            #print(row[j].text)
        print(i)

    curr.close()
    conn.close()
except Exception as e:
    doc.add_paragraph(str(e))

#based on abetreportgenerator.py shenanigans
out_dir = "generated_pdfs"
base = "report"
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, f"{base}_ABET_Report.docx")
doc.save(out_path)