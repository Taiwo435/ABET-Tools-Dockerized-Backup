from report.data.appendix_c_data import get_data
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def build(questionnaire):
    """
    Build template context for Appendix C - Equipment.

    The returned keys are duplicated at the top level and nested under
    ``equipment`` so the DOCX template can use either style.
    """
    equipment = get_data(questionnaire)

    context = {
        "equipment": equipment,
        **equipment,
    }

    return context


def _set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    for attr in ("w:themeFill", "w:themeFillTint", "w:themeFillShade"):
        qualified = qn(attr)
        if shd.get(qualified) is not None:
            del shd.attrib[qualified]


def apply_post_render_formatting(output_path: str):
    """
    Restore Appendix C row styling that cannot be expressed cleanly with a
    single repeating template row.

    The second "PC Workstations / Quantity" table in Appendix C is the SCIA
    instructional labs table. Room rows in the source questionnaire use a gray
    fill, so we apply that styling after docxtpl has expanded the table rows.
    """
    doc = Document(output_path)

    pc_workstation_tables_seen = 0

    for table in doc.tables:
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if header[:2] != ["PC Workstations", "Quantity"]:
            continue

        pc_workstation_tables_seen += 1
        if pc_workstation_tables_seen == 2:
            for cell in table.rows[0].cells:
                _set_cell_fill(cell, "C1E4F5")

        if pc_workstation_tables_seen != 2:
            continue

        for row in table.rows[1:]:
            label = row.cells[0].text.strip()
            quantity = row.cells[1].text.strip()
            if label and not quantity and not label.startswith("{"):
                for cell in row.cells:
                    _set_cell_fill(cell, "D9D9D9")

    for table in doc.tables:
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if header[:1] != ["Printers"]:
            continue

        for cell in table.rows[0].cells:
            _set_cell_fill(cell, "C1E4F5")
        break

    doc.save(output_path)
